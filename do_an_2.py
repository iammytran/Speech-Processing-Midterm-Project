from do_an_1 import *

from docx import Document
import re
from collections import defaultdict

dict_file = "VDic_uni.docx"
doc = Document(dict_file)

# 1. read file doc line-by-line, đọc từng chữ
words = []

def process_dict(path):
    doc = Document(path)
    paras = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts = text.split("\t")
            words.append(parts[0])
        paras.append(text)
    return paras


paragraphs = process_dict(dict_file)
print(f"Từ điển có: {len(paragraphs)}")
# print(paragraphs[:5])   # list of strings
# print(words[:3])

# 2. tách ra nếu cần thiết 
syllables = []
for word in words:
    syllable_split = re.split(r"[\s\-]+", word)
    for syllable in syllable_split:
        syllables.append(syllable)
print(f"Có {len(syllables)} âm tiết")

# 3. process từng âm tiết -> nếu process dc thì thêm nó vào list [âm tiết TV], 0 thì thêm vào list [ko phải âm tiết TV]
valid_syllables = defaultdict(int)
invalid_syllables = defaultdict(int)

for s in syllables:
    try:
        ipa = VietnameseSyllable(s).to_ipa()
        valid_syllables[s] +=1
    except (ValueError, InvalidSyllableError) as e:  # 
        # print(f"[WARN] Cannot convert: {word} ({e})")
        invalid_syllables[s]+=1

print(f"Số lượng âm tiết tiếng việt: {len(valid_syllables)}")
print(f"Số lượng âm tiết không phải tiếng việt: {len(invalid_syllables)}")